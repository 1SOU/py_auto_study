import xlwings as xw
from docx import Document

if __name__ == '__main__':
    文件= Document('1.docx')
    # print(文件.paragraphs[0]) # 段落对象，并不是文本
    # 段落= 文件.paragraphs[0]
    # 块 = 段落.runs # 段落中的每一块
    # for 文字 in 块:
    #     print(文字.text)

    count =0
    for 段落 in 文件.paragraphs:
        print(段落.text)
        if '我' in 段落.text:
            count +=1

    print(count)
