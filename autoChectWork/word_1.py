



from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.text import 


doc = Document('2.docx') # 若是不写文件名字，就新建

# doc.add_heading('一级标题',level=1)
# doc.add_paragraph('正文1') # 可以直接加正文
# # 也可以 按块的 分格式加正文
# runs=doc.add_paragraph('正文1')
# runs.add_run('aaa').bold= True
# runs.add_run('aaa').italic= True
# runs.add_run('aaa') # 正常

runs=doc.add_paragraph('正文1')
runs.alignement= WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save('2.docx')
